"""Text extraction, split by page, with geometry where the format carries it.

Page granularity is the point: a verified fact cites a document *and a page*, so
"where did this figure come from" is answerable down to the page a reviewer can
open. Native-text PDFs go one better and yield word rectangles, which is what
lets the reviewer see the passage highlighted on the original page instead of on
a transcript of it. Formats we cannot read are marked NEEDS_OCR rather than
silently stored as empty text.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field

from ..domain.enums import DocumentStatus
from ..provenance.geometry import NONE, TEXT, PageGeometry, Word
from . import pdf_geometry


@dataclass(frozen=True)
class ExtractedPage:
    """One page of canonical text, plus geometry when the format provides it.

    ``words`` being empty is not a failure: a ``.txt`` file has characters but
    no rendered layout, so a citation into it is exact at the *span* level and
    silent about rectangles. The distinction is preserved all the way to the UI.
    """

    text: str
    width: float | None = None
    height: float | None = None
    extraction_method: str = TEXT
    words: tuple[Word, ...] = field(default=())

    @property
    def has_geometry(self) -> bool:
        return bool(self.words)


@dataclass(frozen=True)
class ExtractionResult:
    pages: list[ExtractedPage]
    status: DocumentStatus
    note: str | None = None

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def page_texts(self) -> list[str]:
        return [page.text for page in self.pages]

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.page_texts)


def _text_pages(texts: list[str], method: str = TEXT) -> list[ExtractedPage]:
    return [ExtractedPage(text=text, extraction_method=method if text else NONE) for text in texts]


def extract(data: bytes, mime_type: str, filename: str) -> ExtractionResult:
    if mime_type in ("text/plain", "text/markdown"):
        return _extract_text(data)
    if mime_type == "application/pdf":
        return _extract_pdf(data)
    if mime_type.startswith("application/vnd.openxmlformats-officedocument.wordprocessingml"):
        return _extract_docx(data)
    if mime_type.startswith("image/"):
        return ExtractionResult(
            pages=_text_pages([""]),
            status=DocumentStatus.NEEDS_OCR,
            note="image upload; OCR not configured, no text indexed",
        )
    return ExtractionResult(
        pages=_text_pages([""]),
        status=DocumentStatus.EXTRACTION_FAILED,
        note=f"no extractor for {mime_type}",
    )


def _extract_text(data: bytes) -> ExtractionResult:
    text = data.decode("utf-8", errors="replace")
    # Honour explicit form feeds; otherwise the whole file is one page.
    pages = text.split("\f") if "\f" in text else [text]
    return ExtractionResult(
        pages=_text_pages([p.strip() for p in pages]), status=DocumentStatus.EXTRACTED
    )


def _extract_pdf(data: bytes) -> ExtractionResult:
    """Geometry-bearing extraction first; plain text only as a fallback."""
    native = _extract_pdf_with_geometry(data)
    if native is not None:
        return native
    return _extract_pdf_text_only(data)


def _extract_pdf_with_geometry(data: bytes) -> ExtractionResult | None:
    pages = pdf_geometry.extract(data)
    if pages is None:
        return None
    if not any(page.text.strip() for page in pages):
        # A scan: the file has pages but no text layer. Say so rather than
        # storing empty strings that look like successfully extracted nothing.
        return ExtractionResult(
            pages=[_from_geometry(page.text, page.geometry) for page in pages] or _text_pages([""]),
            status=DocumentStatus.NEEDS_OCR,
            note="PDF contains no extractable text layer (likely a scan)",
        )
    return ExtractionResult(
        pages=[_from_geometry(page.text, page.geometry) for page in pages],
        status=DocumentStatus.EXTRACTED,
    )


def _from_geometry(text: str, geometry: PageGeometry) -> ExtractedPage:
    return ExtractedPage(
        text=text,
        width=geometry.width,
        height=geometry.height,
        extraction_method=geometry.extraction_method if geometry.words else NONE,
        words=geometry.words,
    )


def _extract_pdf_text_only(data: bytes) -> ExtractionResult:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return ExtractionResult(
            pages=_text_pages([""]),
            status=DocumentStatus.NEEDS_OCR,
            note="no PDF text backend installed; PDF stored without text extraction",
        )
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:  # pragma: no cover - depends on the file
        return ExtractionResult(
            pages=_text_pages([""]),
            status=DocumentStatus.EXTRACTION_FAILED,
            note=f"pypdf error: {exc}",
        )
    if not any(pages):
        return ExtractionResult(
            pages=_text_pages(pages or [""]),
            status=DocumentStatus.NEEDS_OCR,
            note="PDF contains no extractable text layer (likely a scan)",
        )
    return ExtractionResult(
        pages=_text_pages(pages),
        status=DocumentStatus.EXTRACTED,
        note="text extracted without page geometry; citations are span-level only",
    )


def _extract_docx(data: bytes) -> ExtractionResult:
    try:
        import docx  # type: ignore
    except ImportError:  # pragma: no cover - dependency is declared
        return ExtractionResult(
            pages=_text_pages([""]),
            status=DocumentStatus.EXTRACTION_FAILED,
            note="python-docx not installed",
        )
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # pragma: no cover - depends on the file
        return ExtractionResult(
            pages=_text_pages([""]),
            status=DocumentStatus.EXTRACTION_FAILED,
            note=f"python-docx error: {exc}",
        )
    blocks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    text = "\n".join(blocks)
    return ExtractionResult(
        pages=_text_pages([text] if text else [""]),
        status=DocumentStatus.EXTRACTED if text else DocumentStatus.NEEDS_OCR,
        note=None if text else "DOCX contained no text",
    )
