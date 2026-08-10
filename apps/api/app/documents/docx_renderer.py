"""DOCX rendering of a demand draft.

The rendered file is a projection of the stored sections — it adds formatting,
never content. Anything that appears in the DOCX came from a section body that
the validation engine has already seen.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ..domain.enums import DocumentType
from ..domain.models import Demand
from ..generation.context import DemandContext
from ..ingestion.storage import ObjectStore

# Sections whose bodies are pre-formatted blocks rather than prose paragraphs.
_PREFORMATTED = {"header", "claim_metadata", "medical_expense_summary", "future_medical", "signature"}
_UNTITLED = {"header", "claim_metadata", "demand_title", "signature"}
_IMAGE_MIME = {"image/png", "image/jpeg"}


def render_docx(
    demand: Demand,
    context: DemandContext,
    *,
    store: ObjectStore | None = None,
    watermark_draft: bool = True,
) -> bytes:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    footer = section.footer.paragraphs[0]
    claim_number = context.claim.claim_number if context.claim else "no claim number"
    footer.text = f"{context.client_name} — Claim {claim_number} — page "
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if watermark_draft and not demand.locked:
        banner = document.add_paragraph()
        run = banner.add_run(f"DRAFT v{demand.version} — NOT FOR RELEASE")
        run.bold = True
        run.font.size = Pt(9)
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for demand_section in sorted(demand.sections, key=lambda s: s.position):
        key = demand_section.key

        if key == "demand_title":
            for line_index, line in enumerate(demand_section.body.splitlines()):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(line)
                run.bold = True
                if line_index == 0:
                    run.underline = True
            continue

        if key not in _UNTITLED:
            heading = document.add_paragraph()
            run = heading.add_run(demand_section.title.upper())
            run.bold = True

        if key in _PREFORMATTED:
            for line in demand_section.body.splitlines():
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
        elif key == "conditions":
            lines = demand_section.body.splitlines()
            if lines:
                document.add_paragraph(lines[0])
            for line in lines[1:]:
                text = line.strip()
                if not text:
                    continue
                text = text.split(". ", 1)[-1] if text[0].isdigit() else text
                document.add_paragraph(text, style="List Number")
        else:
            for block in demand_section.body.split("\n\n"):
                if block.strip():
                    document.add_paragraph(block.strip())

        if key == "photographs" and store is not None:
            _embed_photographs(document, context, store)

        document.add_paragraph()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _embed_photographs(document, context: DemandContext, store: ObjectStore) -> None:
    for photo in context.documents_of_type(DocumentType.PHOTOGRAPH):
        if photo.mime_type not in _IMAGE_MIME:
            continue
        try:
            data = store.get(photo.storage_key)
            document.add_picture(io.BytesIO(data), width=Inches(4.5))
        except Exception:  # pragma: no cover - a bad image must not break the letter
            document.add_paragraph(f"[Photograph {photo.original_filename} could not be embedded]")
