"""Binding writes into the attorney's own file and nothing else."""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import golden
from app.templates import analyzer, binder

VALUES = {
    "letter_date": "June 6, 2025",
    "adjuster_name": "T. Nakamura",
    "carrier_name": "Meridian Casualty Insurance",
    "client_name": "Patrick Donahue",
    "claim_number": "017204635",
    "insured_name": "Marisol Reyes",
    "date_of_loss": "March 4, 2024",
    "policy_number": "CA-88213",
    "demand_expiration": "July 6, 2025 at 5:00 PM",
    "introduction_section": ["First paragraph.", "Second paragraph."],
    "liability_section": ["The insured driver struck our client from behind."],
    "medical_treatment_section": ["Treatment prose."],
    "imaging_section": ["Imaging prose."],
    "future_medical_section": ["Future care line one.\nFuture care line two."],
    "medical_expenses": [
        {"provider": "Vermont Spine and Injury", "description": "chiropractic", "amount": "$6,480.00"},
        {"provider": "MAX MRI Radiology", "description": "MRI lumbar", "amount": "$3,500.00"},
        {"provider": "Harbor Pain Management", "description": "injection", "amount": "Pending"},
    ],
    "medical_expenses_total": "$9,980.00",
    "pain_and_suffering_section": ["Pain prose."],
    "demand_section": ["Demand prose."],
    "conditions_section": ["Condition one.", "Condition two."],
    "attorney_name": "Dana Okafor",
    "firm_name": "Stalwart Law Group",
}


@pytest.fixture(scope="module")
def template_bytes() -> bytes:
    return golden.TEMPLATE_PATH.read_bytes()


@pytest.fixture(scope="module")
def manifest(template_bytes):
    return analyzer.analyze(template_bytes)


@pytest.fixture(scope="module")
def bound(template_bytes, manifest):
    return binder.bind(template_bytes, manifest, VALUES)


def _document(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def test_every_placeholder_is_gone(bound):
    data, _ = bound
    assert binder.unbound_placeholders(data) == []


def test_inline_values_land_in_their_table_cells(bound):
    data, _ = bound
    rows = [[cell.text for cell in row.cells] for row in _document(data).tables[0].rows]
    assert ["Our Client", "Patrick Donahue"] in rows
    assert ["Claim Number", "017204635"] in rows
    assert ["Policy Number", "CA-88213"] in rows


def test_a_block_slot_becomes_one_paragraph_per_item(bound):
    data, _ = bound
    texts = [p.text for p in _document(data).paragraphs]
    assert "First paragraph." in texts
    assert "Second paragraph." in texts


def test_a_block_slot_keeps_the_placeholder_paragraph_formatting(bound):
    data, _ = bound
    # The letter date placeholder was right-aligned in the template.
    paragraph = next(p for p in _document(data).paragraphs if p.text == "June 6, 2025")
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_newlines_inside_a_block_become_real_line_breaks(bound):
    data, _ = bound
    document = _document(data)
    paragraph = next(
        p for p in document.paragraphs if "Future care line one." in p.text
    )
    assert len(paragraph._p.findall(".//" + qn("w:br"))) == 1
    assert "Future care line two." in paragraph.text


def test_a_row_slot_repeats_the_sample_row(bound):
    data, _ = bound
    table = _document(data).tables[1]
    assert [cell.text for cell in table.rows[0].cells] == ["Provider", "Description", "Amount"]
    assert len(table.rows) == 4  # header + three bills
    assert [cell.text for cell in table.rows[3].cells][2] == "Pending"


def test_a_repeated_row_keeps_the_sample_row_formatting(template_bytes, manifest, bound):
    data, _ = bound
    original = _document(template_bytes).tables[1].rows[1]
    generated = _document(data).tables[1].rows[1]
    assert original._tr.find(qn("w:trPr")) is None or generated._tr.find(qn("w:trPr")) is not None
    # Cell widths carry the table's column geometry; they must be identical.
    original_widths = [c._tc.find(qn("w:tcPr")).find(qn("w:tcW")).get(qn("w:w"))
                       for c in original.cells]
    generated_widths = [c._tc.find(qn("w:tcPr")).find(qn("w:tcW")).get(qn("w:w"))
                        for c in generated.cells]
    assert original_widths == generated_widths


def test_the_report_counts_what_was_written(bound, manifest):
    _, report = bound
    assert report.template_sha256 == manifest.fingerprint.sha256
    assert report.table_rows == 3
    assert report.inline_replacements >= 6
    assert report.block_paragraphs >= 10
    assert set(report.bound_slots) == set(manifest.slot_names())


def test_an_unbound_slot_is_an_error_not_an_empty_string(template_bytes, manifest):
    partial = {k: v for k, v in VALUES.items() if k != "claim_number"}
    with pytest.raises(binder.UnboundSlotError) as exc:
        binder.bind(template_bytes, manifest, partial)
    assert "claim_number" in exc.value.names


def test_a_slot_may_be_skipped_only_when_the_caller_says_so(template_bytes, manifest):
    partial = {k: v for k, v in VALUES.items() if k != "policy_number"}
    data, report = binder.bind(
        template_bytes, manifest, partial, allow_missing=["policy_number"]
    )
    assert "policy_number" not in report.bound_slots
    # The placeholder is still there — the caller took responsibility for it.
    assert binder.unbound_placeholders(data) == ["policy_number"]


def test_wrong_value_shape_is_rejected(template_bytes, manifest):
    bad = dict(VALUES, medical_expenses="not a list of rows")
    with pytest.raises(binder.SlotBindingError):
        binder.bind(template_bytes, manifest, bad)


def test_binding_does_not_mutate_the_source_bytes(template_bytes, manifest):
    before = bytes(template_bytes)
    binder.bind(template_bytes, manifest, VALUES)
    assert template_bytes == before
