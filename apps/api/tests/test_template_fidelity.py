"""Fidelity validation catches every way a bound document can drift.

Each negative test deliberately damages the generated document in one specific
way and asserts the matching TEMPLATE_00x code fires. A rule nobody can make
fail is a rule nobody should trust.
"""

from __future__ import annotations

import copy
import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import golden
from app.templates import analyzer, binder, fidelity
from test_template_binder import VALUES


@pytest.fixture(scope="module")
def template_bytes() -> bytes:
    return golden.TEMPLATE_PATH.read_bytes()


@pytest.fixture(scope="module")
def manifest(template_bytes):
    return analyzer.analyze(template_bytes)


@pytest.fixture(scope="module")
def bound_bytes(template_bytes, manifest) -> bytes:
    data, _ = binder.bind(template_bytes, manifest, VALUES)
    return data


def _resave(document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _codes(report) -> set[str]:
    return {issue.code for issue in report.issues}


# --------------------------------------------------------------------- happy path


def test_a_faithfully_bound_document_passes(template_bytes, manifest, bound_bytes):
    report = fidelity.verify(template_bytes, bound_bytes, manifest)
    assert report.is_faithful
    assert report.issues == ()
    assert report.required_blocks_preserved == report.required_blocks_expected
    assert report.styles_changed == 0
    assert report.headers_changed == 0
    assert report.footers_changed == 0
    assert report.page_setup_changed is False


def test_the_report_serializes_in_the_documented_shape(template_bytes, manifest, bound_bytes):
    payload = fidelity.verify(template_bytes, bound_bytes, manifest).to_dict()
    assert payload["template_hash"] == manifest.fingerprint.sha256
    assert set(payload["required_blocks"]) == {"expected", "preserved"}
    assert payload["blocking_issues"] == []


# ------------------------------------------------------------------ negative cases


def test_template_001_fires_when_sections_are_reordered(template_bytes, manifest, bound_bytes):
    document = Document(io.BytesIO(bound_bytes))
    body = document.element.body
    headings = [
        p for p in body.findall(qn("w:p"))
        if "".join(t.text or "" for t in p.iter(qn("w:t"))) in ("LIABILITY", "MEDICAL TREATMENT")
    ]
    assert len(headings) == 2
    first, second = headings
    second.addnext(first)  # swap the order the letter argues in

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_SECTION_ORDER in _codes(report)
    assert not report.is_faithful


def test_template_002_fires_when_a_required_block_is_deleted(
    template_bytes, manifest, bound_bytes
):
    document = Document(io.BytesIO(bound_bytes))
    heading = next(
        p for p in document.paragraphs if p.text == "CONDITIONS OF ACCEPTANCE"
    )
    heading._p.getparent().remove(heading._p)

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_BLOCK_MISSING in _codes(report)


def test_template_003_fires_when_the_footer_changes(template_bytes, manifest, bound_bytes):
    document = Document(io.BytesIO(bound_bytes))
    document.sections[0].footer.paragraphs[0].add_run(" — DRAFT, DO NOT SEND")

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_HEADER_FOOTER in _codes(report)
    assert report.footers_changed == 1


def test_template_003_fires_when_the_header_changes(template_bytes, manifest, bound_bytes):
    document = Document(io.BytesIO(bound_bytes))
    document.sections[0].header.paragraphs[0].add_run(" (rebranded)")

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_HEADER_FOOTER in _codes(report)
    assert report.headers_changed == 1


def test_template_004_fires_when_page_setup_changes(template_bytes, manifest, bound_bytes):
    document = Document(io.BytesIO(bound_bytes))
    document.sections[0].left_margin = Inches(2)

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_PAGE_SETUP in _codes(report)
    assert report.page_setup_changed is True


def test_template_005_fires_when_a_protected_style_changes(
    template_bytes, manifest, bound_bytes
):
    document = Document(io.BytesIO(bound_bytes))
    document.styles["Normal"].font.size = Pt(18)

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_STYLES in _codes(report)
    assert report.styles_changed == 1


def test_template_006_fires_when_a_protected_table_changes(template_bytes, manifest):
    """A table with no ROW slot must keep its exact shape."""
    document = Document(io.BytesIO(template_bytes))
    static = document.add_table(rows=2, cols=2)
    static.style = "Table Grid"
    static.rows[0].cells[0].text = "Enclosures"
    static.rows[0].cells[1].text = "Count"
    static.rows[1].cells[0].text = "Photographs"
    static.rows[1].cells[1].text = "12"
    with_static = _resave(document)
    static_manifest = analyzer.analyze(with_static)

    bound, _ = binder.bind(with_static, static_manifest, VALUES)
    assert fidelity.verify(with_static, bound, static_manifest).is_faithful

    damaged = Document(io.BytesIO(bound))
    damaged.tables[-1].add_row().cells[0].text = "Smuggled row"
    report = fidelity.verify(with_static, _resave(damaged), static_manifest)
    assert fidelity.TEMPLATE_TABLE in _codes(report)


def test_template_007_warns_when_pagination_differs(template_bytes, manifest, bound_bytes):
    report = fidelity.compare(
        manifest, analyzer.analyze(bound_bytes), template_pages=3, generated_pages=5
    )
    pagination = [i for i in report.issues if i.code == fidelity.TEMPLATE_PAGINATION]
    assert len(pagination) == 1
    assert pagination[0].severity == fidelity.WARNING
    # A pagination difference is worth a look; it is not a reason to block.
    assert report.is_faithful


def test_template_008_fires_when_immutable_text_is_edited(
    template_bytes, manifest, bound_bytes
):
    document = Document(io.BytesIO(bound_bytes))
    heading = next(p for p in document.paragraphs if p.text == "LIABILITY")
    heading.runs[0].text = "FAULT"

    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert fidelity.TEMPLATE_OOXML_BLOCK in _codes(report)


def test_a_pristine_copy_of_the_template_compares_clean(template_bytes, manifest):
    """Round-tripping through python-docx must not look like a mutation."""
    document = Document(io.BytesIO(template_bytes))
    report = fidelity.verify(template_bytes, _resave(document), manifest)
    assert report.is_faithful, [i.message for i in report.blocking_issues]


def test_comparison_does_not_depend_on_xml_serialization_details(template_bytes, manifest):
    """Attribute order and self-closing tags are not template changes."""
    original = analyzer.package_part_digests(template_bytes)
    document = Document(io.BytesIO(template_bytes))
    resaved = analyzer.package_part_digests(_resave(document))
    assert original["word/styles.xml"] == resaved["word/styles.xml"]
    assert original["word/header1.xml"] == resaved["word/header1.xml"]
    assert original["word/footer1.xml"] == resaved["word/footer1.xml"]


def test_manifest_comparison_is_pure(manifest, bound_bytes):
    """compare() must not mutate either manifest it is handed."""
    generated = analyzer.analyze(bound_bytes)
    before_left = copy.deepcopy(manifest.to_dict())
    before_right = copy.deepcopy(generated.to_dict())
    fidelity.compare(manifest, generated)
    assert manifest.to_dict() == before_left
    assert generated.to_dict() == before_right
