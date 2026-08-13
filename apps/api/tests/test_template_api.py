"""HTTP surface for template upload, inspection, and binding."""

from __future__ import annotations

import io

import pytest
from docx import Document

import golden
from conftest import ATTORNEY, READONLY

DOCX_MIME = golden.DOCX_MIME


def _upload(client, case_id, data=None, filename="template.docx", headers=ATTORNEY):
    payload = golden.TEMPLATE_PATH.read_bytes() if data is None else data
    return client.post(
        f"/v1/cases/{case_id}/templates",
        files={"file": (filename, payload, DOCX_MIME)},
        data={"name": "Firm standard"},
        headers=headers,
    )


def test_upload_returns_the_analyzed_manifest(client, seeded_case):
    response = _upload(client, seeded_case["case_id"])
    assert response.status_code == 201, response.text
    body = response.json()

    assert body["sha256"]
    assert body["block_count"] > 20
    assert "liability_section" in body["slot_names"]
    assert body["header_parts"] == ["word/header1.xml"]
    assert body["footer_parts"] == ["word/footer1.xml"]
    assert body["page_setup"]["page_width"] == 12240
    assert body["unknown_slots"] == []

    kinds = {slot["name"]: slot["kind"] for slot in body["slots"]}
    assert kinds["medical_expenses"] == "row"
    assert kinds["liability_section"] == "block"
    assert kinds["client_name"] == "inline"


def test_a_file_that_is_not_a_docx_is_rejected(client, seeded_case):
    response = _upload(client, seeded_case["case_id"], data=b"plain text", filename="notes.txt")
    assert response.status_code == 400
    assert "template" in response.json()["detail"].lower()


def test_an_empty_upload_is_rejected(client, seeded_case):
    response = _upload(client, seeded_case["case_id"], data=b"")
    assert response.status_code == 400


def test_uploading_the_same_template_twice_conflicts(client, seeded_case):
    case_id = seeded_case["case_id"]
    first = _upload(client, case_id)
    assert first.status_code == 201
    second = _upload(client, case_id)
    assert second.status_code == 409
    assert second.json()["detail"]["existing_template_id"] == first.json()["id"]


def test_a_readonly_user_cannot_upload_a_template(client, seeded_case):
    response = _upload(client, seeded_case["case_id"], headers=READONLY)
    assert response.status_code == 403


def test_a_readonly_user_can_read_templates(client, seeded_case):
    case_id = seeded_case["case_id"]
    _upload(client, case_id)
    response = client.get(f"/v1/cases/{case_id}/templates", headers=READONLY)
    assert response.status_code == 200
    assert len(response.json()) == 1


def test_a_template_using_an_unknown_slot_is_reported_not_silently_accepted(
    client, seeded_case
):
    document = Document(io.BytesIO(golden.TEMPLATE_PATH.read_bytes()))
    document.add_paragraph("{{punitive_multiplier}}")
    buffer = io.BytesIO()
    document.save(buffer)

    response = _upload(client, seeded_case["case_id"], data=buffer.getvalue())
    assert response.status_code == 201
    body = response.json()
    assert body["unknown_slots"] == ["punitive_multiplier"]
    assert {s["name"]: s["resolvable"] for s in body["slots"]}["punitive_multiplier"] is False


def test_generation_fails_loudly_when_a_template_slot_has_no_resolver(client, seeded_case):
    """A template the system cannot fill must not produce a half-filled letter."""
    case_id = seeded_case["case_id"]
    document = Document(io.BytesIO(golden.TEMPLATE_PATH.read_bytes()))
    document.add_paragraph("{{punitive_multiplier}}")
    buffer = io.BytesIO()
    document.save(buffer)
    template = _upload(client, case_id, data=buffer.getvalue()).json()

    demand = client.post(f"/v1/cases/{case_id}/demands", json={}, headers=ATTORNEY).json()
    client.post(
        f"/v1/demands/{demand['id']}/template",
        json={"template_id": template["id"]},
        headers=ATTORNEY,
    )
    client.post(f"/v1/demands/{demand['id']}/generate", json={}, headers=ATTORNEY)

    issues = client.post(
        f"/v1/demands/{demand['id']}/validate", headers=ATTORNEY
    ).json()
    codes = {issue["code"] for issue in issues if issue["severity"] == "BLOCKING"}
    assert "TEMPLATE_002" in codes

    response = client.get(f"/v1/demands/{demand['id']}/docx", headers=ATTORNEY)
    # The binder refuses outright; a half-filled letter is never served.
    assert response.status_code == 409
    assert "punitive_multiplier" in response.json()["detail"]


def test_binding_a_template_to_a_demand_records_the_hash(client, seeded_case_with_facts):
    case_id = seeded_case_with_facts["case_id"]
    template = _upload(client, case_id).json()
    demand = client.post(f"/v1/cases/{case_id}/demands", json={}, headers=ATTORNEY).json()

    response = client.post(
        f"/v1/demands/{demand['id']}/template",
        json={"template_id": template["id"]},
        headers=ATTORNEY,
    )
    assert response.status_code == 200

    detail = client.get(f"/v1/demands/{demand['id']}", headers=ATTORNEY).json()
    assert detail["template_id"] == template["id"]
    assert detail["template_sha256"] == template["sha256"]


def test_binding_an_unknown_template_is_a_404(client, seeded_case):
    case_id = seeded_case["case_id"]
    demand = client.post(f"/v1/cases/{case_id}/demands", json={}, headers=ATTORNEY).json()
    response = client.post(
        f"/v1/demands/{demand['id']}/template",
        json={"template_id": "tpl_does_not_exist"},
        headers=ATTORNEY,
    )
    assert response.status_code == 404


def test_a_template_from_another_case_cannot_be_bound(client, seeded_case):
    other = client.post(
        "/v1/cases",
        json={"reference": "OTHER-1", "client_display_name": "Someone Else"},
        headers=ATTORNEY,
    ).json()
    template = _upload(client, other["id"]).json()

    case_id = seeded_case["case_id"]
    demand = client.post(f"/v1/cases/{case_id}/demands", json={}, headers=ATTORNEY).json()
    response = client.post(
        f"/v1/demands/{demand['id']}/template",
        json={"template_id": template["id"]},
        headers=ATTORNEY,
    )
    assert response.status_code == 409


def test_fidelity_report_is_unavailable_until_the_demand_is_validated(
    client, seeded_case_with_facts
):
    case_id = seeded_case_with_facts["case_id"]
    demand = client.post(f"/v1/cases/{case_id}/demands", json={}, headers=ATTORNEY).json()
    response = client.get(f"/v1/demands/{demand['id']}/fidelity", headers=ATTORNEY)
    assert response.status_code == 409
    assert "no template" in response.json()["detail"]


def test_fidelity_report_is_served_after_validation(client, seeded_case_with_facts):
    case_id = seeded_case_with_facts["case_id"]
    template = _upload(client, case_id).json()
    demand = client.post(f"/v1/cases/{case_id}/demands", json={}, headers=ATTORNEY).json()
    client.post(
        f"/v1/demands/{demand['id']}/template",
        json={"template_id": template["id"]},
        headers=ATTORNEY,
    )
    client.post(f"/v1/demands/{demand['id']}/generate", json={}, headers=ATTORNEY)
    client.post(f"/v1/demands/{demand['id']}/validate", headers=ATTORNEY)

    response = client.get(f"/v1/demands/{demand['id']}/fidelity", headers=ATTORNEY)
    assert response.status_code == 200
    report = response.json()
    assert report["template_hash"] == template["sha256"]
    assert report["blocking_issues"] == []
    assert report["required_blocks"]["preserved"] == report["required_blocks"]["expected"]


def test_the_known_slot_catalog_is_published(client):
    response = client.get("/v1/template-slots", headers=ATTORNEY)
    assert response.status_code == 200
    names = response.json()
    assert "client_name" in names
    assert "medical_expenses" in names
    assert "liability_section" in names


def test_template_upload_is_recorded_in_the_audit_trail(client, seeded_case):
    case_id = seeded_case["case_id"]
    template = _upload(client, case_id).json()

    events = client.get(f"/v1/cases/{case_id}/audit", headers=ATTORNEY).json()
    ingested = [e for e in events if e["event"] == "TEMPLATE_INGESTED"]
    assert len(ingested) == 1
    assert ingested[0]["payload"]["sha256"] == template["sha256"]
    assert "liability_section" in ingested[0]["payload"]["slots"]


@pytest.mark.invariant
def test_the_stored_template_bytes_are_immutable(client, seeded_case, db):
    """INVARIANT-006 — the substrate the letter is built on cannot be swapped."""
    from app.domain.models import LetterTemplate
    from app.ingestion.storage import ImmutableObjectError, get_object_store

    case_id = seeded_case["case_id"]
    template_id = _upload(client, case_id).json()["id"]
    template = db.get(LetterTemplate, template_id)

    with pytest.raises(ImmutableObjectError):
        get_object_store().put(template.storage_key, b"a different template", immutable=True)
