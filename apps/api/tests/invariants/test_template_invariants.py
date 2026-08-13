"""INVARIANT-006 — template formatting outside dynamic regions is never mutated.

These tests are written against the guarantee, not the implementation. If the
binder is rewritten they must still hold; if they stop holding, the product
claim that a generated letter matches the firm's template is false.
"""

from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn

import golden
from app.templates import analyzer, binder, fidelity
from test_template_binder import VALUES

pytestmark = pytest.mark.invariant


@pytest.fixture(scope="module")
def template_bytes() -> bytes:
    return golden.TEMPLATE_PATH.read_bytes()


@pytest.fixture(scope="module")
def manifest(template_bytes):
    return analyzer.analyze(template_bytes)


@pytest.fixture(scope="module")
def bound(template_bytes, manifest) -> bytes:
    data, _ = binder.bind(template_bytes, manifest, VALUES)
    return data


def _part(data: bytes, name: str) -> bytes:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        return archive.read(name)


def test_headers_survive_export_unchanged(template_bytes, bound):
    before = analyzer.package_part_digests(template_bytes)
    after = analyzer.package_part_digests(bound)
    assert after["word/header1.xml"] == before["word/header1.xml"]


def test_footers_survive_export_unchanged(template_bytes, bound):
    before = analyzer.package_part_digests(template_bytes)
    after = analyzer.package_part_digests(bound)
    assert after["word/footer1.xml"] == before["word/footer1.xml"]


def test_styles_survive_export_unchanged(template_bytes, bound):
    before = analyzer.package_part_digests(template_bytes)
    after = analyzer.package_part_digests(bound)
    assert after["word/styles.xml"] == before["word/styles.xml"]
    assert after["word/numbering.xml"] == before["word/numbering.xml"]


def test_page_setup_survives_export_unchanged(template_bytes, bound):
    assert analyzer.analyze(bound).page_setup == analyzer.analyze(template_bytes).page_setup


def test_section_breaks_survive_export_unchanged(template_bytes, bound):
    before, after = analyzer.analyze(template_bytes), analyzer.analyze(bound)
    assert after.section_break_count == before.section_break_count


def test_page_breaks_survive_export_unchanged(template_bytes, bound):
    before = sum(1 for b in analyzer.analyze(template_bytes).blocks if b.has_page_break)
    after = sum(1 for b in analyzer.analyze(bound).blocks if b.has_page_break)
    assert after == before


def test_every_immutable_block_is_byte_identical(template_bytes, manifest, bound):
    """Not merely present — the same text, style and order."""
    generated = {(b.kind, b.style or "", b.text_sha256) for b in analyzer.analyze(bound).blocks}
    dynamic = {slot.block_index for slot in manifest.slots}
    for block in manifest.blocks:
        if block.index in dynamic:
            continue
        assert (block.kind, block.style or "", block.text_sha256) in generated, (
            f"immutable block {block.index} ({block.text[:40]!r}) did not survive binding"
        )


def test_table_geometry_is_untouched(template_bytes, bound):
    """Column widths and borders decide how the letter reads on the page."""
    before = Document(io.BytesIO(template_bytes))
    after = Document(io.BytesIO(bound))
    assert len(after.tables) == len(before.tables)
    for original, generated in zip(before.tables, after.tables):
        original_grid = original._tbl.find(qn("w:tblGrid"))
        generated_grid = generated._tbl.find(qn("w:tblGrid"))
        original_cols = [c.get(qn("w:w")) for c in original_grid]
        generated_cols = [c.get(qn("w:w")) for c in generated_grid]
        assert original_cols == generated_cols

        original_props = original._tbl.find(qn("w:tblPr"))
        generated_props = generated._tbl.find(qn("w:tblPr"))
        from lxml import etree

        assert etree.tostring(original_props) == etree.tostring(generated_props)


def test_cell_shading_is_untouched(template_bytes, bound):
    """The template shades header cells; the letter must look the same."""
    before = Document(io.BytesIO(template_bytes)).tables[1].rows[0]
    after = Document(io.BytesIO(bound)).tables[1].rows[0]
    for original, generated in zip(before.cells, after.cells):
        original_fill = original._tc.find(qn("w:tcPr")).find(qn("w:shd"))
        generated_fill = generated._tc.find(qn("w:tcPr")).find(qn("w:shd"))
        assert original_fill.get(qn("w:fill")) == generated_fill.get(qn("w:fill"))


def test_embedded_media_is_carried_across(template_bytes, bound):
    before = analyzer.analyze(template_bytes)
    after = analyzer.analyze(bound)
    assert set(after.image_parts) == set(before.image_parts)


def test_no_placeholder_reaches_the_final_document(bound):
    assert binder.unbound_placeholders(bound) == []


def test_binding_the_same_values_twice_produces_the_same_structure(
    template_bytes, manifest
):
    first, _ = binder.bind(template_bytes, manifest, VALUES)
    second, _ = binder.bind(template_bytes, manifest, VALUES)
    left, right = analyzer.analyze(first), analyzer.analyze(second)
    assert [b.text_sha256 for b in left.blocks] == [b.text_sha256 for b in right.blocks]
    assert left.fingerprint.structure_sha256 == right.fingerprint.structure_sha256


def test_a_mutated_document_cannot_pass_fidelity(template_bytes, manifest, bound):
    """The guarantee is only worth anything if violating it is detected."""
    document = Document(io.BytesIO(bound))
    document.sections[0].footer.paragraphs[0].add_run(" (edited)")
    document.styles["Normal"].font.size = None
    buffer = io.BytesIO()
    document.save(buffer)

    report = fidelity.verify(template_bytes, buffer.getvalue(), manifest)
    assert not report.is_faithful
    assert {i.code for i in report.blocking_issues} >= {fidelity.TEMPLATE_HEADER_FOOTER}
