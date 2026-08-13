"""Builds the golden-case fixtures used by the template-fidelity tests.

Run this only when the fixture itself should change:

    python scripts/build_golden_fixture.py

It writes ``apps/api/tests/fixtures/golden_case/template.docx`` — a realistic
attorney demand letter with a header, a footer, heading styles, a numbered
list, two tables and a page break — plus the plain-text case materials that
Phase 3's extraction tests read.

The expected demand document is produced separately by
``scripts/build_golden_expected.py``, which runs the real pipeline. Keeping the
two apart means the committed expectation is the output of the system under
test, not a second hand-written copy of it.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
FIXTURE_DIR = REPO_ROOT / "apps" / "api" / "tests" / "fixtures" / "golden_case"
MATERIALS_DIR = FIXTURE_DIR / "case-materials"

FIRM = "Stalwart Law Group"
FIRM_ADDRESS = "1055 W 7th St, Suite 2800  |  Los Angeles, CA 90017  |  (213) 000-0000"


def _shade(cell, hex_fill: str) -> None:
    shading = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _page_number_field(paragraph) -> None:
    """A real PAGE field, so pagination is Word's job and not ours."""
    run = paragraph.add_run()
    begin = run._r.makeelement(qn("w:fldChar"), {})
    begin.set(qn("w:fldCharType"), "begin")
    instr = run._r.makeelement(qn("w:instrText"), {})
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = run._r.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build_template() -> Document:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    heading = document.styles["Heading 1"]
    heading.font.name = "Georgia"
    heading.font.size = Pt(12)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firm_run = header.add_run(FIRM)
    firm_run.bold = True
    firm_run.font.size = Pt(14)
    header.add_run("\n" + FIRM_ADDRESS).font.size = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CONFIDENTIAL SETTLEMENT COMMUNICATION  —  Page ").font.size = Pt(8)
    _page_number_field(footer)

    # ------------------------------------------------------------- letter head
    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run("{{letter_date}}")

    document.add_paragraph("{{adjuster_name}}")
    document.add_paragraph("{{carrier_name}}")
    document.add_paragraph("Sent via email")
    document.add_paragraph()

    # -------------------------------------------------------- claim metadata
    metadata = document.add_table(rows=5, cols=2)
    metadata.style = "Table Grid"
    rows = [
        ("Our Client", "{{client_name}}"),
        ("Claim Number", "{{claim_number}}"),
        ("Your Insured", "{{insured_name}}"),
        ("Date of Loss", "{{date_of_loss}}"),
        ("Policy Number", "{{policy_number}}"),
    ]
    for row, (label, placeholder) in zip(metadata.rows, rows):
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        _shade(row.cells[0], "EEF1F6")
        row.cells[1].text = placeholder

    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TIME-LIMITED POLICY LIMITS DEMAND")
    title_run.bold = True
    title_run.underline = True

    expiry = document.add_paragraph()
    expiry.alignment = WD_ALIGN_PARAGRAPH.CENTER
    expiry.add_run("This offer expires on {{demand_expiration}}.").italic = True

    # ------------------------------------------------------------- narrative
    narrative_sections = [
        ("INTRODUCTION", "{{introduction_section}}"),
        ("LIABILITY", "{{liability_section}}"),
        ("MEDICAL TREATMENT", "{{medical_treatment_section}}"),
        ("DIAGNOSTIC IMAGING", "{{imaging_section}}"),
        ("FUTURE MEDICAL CARE", "{{future_medical_section}}"),
    ]
    for title_text, placeholder in narrative_sections:
        document.add_paragraph(title_text, style="Heading 1")
        document.add_paragraph(placeholder)

    # -------------------------------------------------------- expenses table
    document.add_paragraph("MEDICAL EXPENSES", style="Heading 1")
    expenses = document.add_table(rows=2, cols=3)
    expenses.style = "Table Grid"
    for index, label in enumerate(("Provider", "Description", "Amount")):
        cell = expenses.rows[0].cells[index]
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        _shade(cell, "1F2A44")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sample = expenses.rows[1].cells
    sample[0].text = "{{medical_expenses[].provider}}"
    sample[1].text = "{{medical_expenses[].description}}"
    sample[2].text = "{{medical_expenses[].amount}}"

    total = document.add_paragraph()
    total.add_run("Total known medical expenses to date: ").bold = True
    total.add_run("{{medical_expenses_total}}")

    document.add_paragraph("PAIN, SUFFERING, AND INCONVENIENCE", style="Heading 1")
    document.add_paragraph("{{pain_and_suffering_section}}")

    document.add_paragraph("DEMAND FOR SETTLEMENT", style="Heading 1")
    document.add_paragraph("{{demand_section}}")

    document.add_paragraph("CONDITIONS OF ACCEPTANCE", style="Heading 1")
    document.add_paragraph("{{conditions_section}}")

    # ------------------------------------------------------------- signature
    document.add_section(WD_SECTION.NEW_PAGE)
    closing = document.add_paragraph("Very truly yours,")
    closing.paragraph_format.space_before = Pt(24)
    document.add_paragraph()
    document.add_paragraph("{{attorney_name}}")
    document.add_paragraph("{{firm_name}}")

    return document


CASE_MATERIALS: dict[str, str] = {
    "police-report.txt": """LOS ANGELES POLICE DEPARTMENT
TRAFFIC COLLISION REPORT

Report Number: TC-2024-118872
Date of Collision: March 4, 2024
Location: Vermont Ave and W 8th St, Los Angeles, CA

PARTY 1 (Driver): Andre Whitfield
  Vehicle: 2019 Toyota Camry, registered to Marisol Reyes
PARTY 2 (Driver): Patrick Donahue
  Vehicle: 2021 Honda Civic

NARRATIVE
Party 2 was stopped at a red signal facing northbound on Vermont Ave.
Party 1, travelling northbound behind Party 2, failed to stop and struck
the rear of Party 2's vehicle. Party 1 stated to the reporting officer
that he "looked down at the navigation screen" before impact.

Primary collision factor: 22350 VC (unsafe speed for conditions).
Party at fault: Party 1.
""",
    "chiropractic-records.txt": """VERMONT SPINE AND INJURY
Patient: Patrick Donahue
Date of Service: March 7, 2024

CHIEF COMPLAINT
Neck and lower back pain beginning after a rear-end motor vehicle
collision on March 4, 2024.

EXAMINATION
Cervical range of motion reduced in extension and right rotation.
Lumbar paraspinal muscle spasm noted bilaterally at L4-S1.

ASSESSMENT
Cervical strain. Lumbar strain with radicular complaints into the right leg.

PLAN
Chiropractic manipulative therapy three times weekly for six weeks.
Re-evaluate in thirty days. Refer for MRI if symptoms persist.
""",
    "mri-report.txt": """MAX MRI RADIOLOGY
Patient: Patrick Donahue
Study Date: December 21, 2024
Examination: MRI LUMBAR SPINE WITHOUT CONTRAST

FINDINGS
L5-S1: Broad-based disc extrusion measuring 9 x 10 x 5 mm, extending
into the right lateral recess with contact upon the traversing right
S1 nerve root.
L4-L5: Mild disc bulge without stenosis.

IMPRESSION
1. L5-S1 disc extrusion with right S1 nerve root contact.
2. Findings correlate with the reported right lower extremity symptoms.
""",
    "billing-summary.txt": """BILLING SUMMARY
Patient: Patrick Donahue

Vermont Spine and Injury ................ $6,480.00   (final)
MAX MRI Radiology ....................... $3,500.00   (final)
Harbor Pain Management .................. NOT YET RECEIVED

Note: the Harbor Pain Management charge for the epidural steroid
injection has not been received. The total above is incomplete.
""",
}


def main() -> None:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    MATERIALS_DIR.mkdir(parents=True, exist_ok=True)

    template_path = FIXTURE_DIR / "template.docx"
    build_template().save(template_path)
    print(f"wrote {template_path.relative_to(REPO_ROOT)}")

    for name, body in CASE_MATERIALS.items():
        path = MATERIALS_DIR / name
        path.write_text(body, encoding="utf-8")
        print(f"wrote {path.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
